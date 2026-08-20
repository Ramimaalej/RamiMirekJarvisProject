"""Create professional PDF and DOCX documents (quality: structured, styled).

Usage (via Jarvis intent `create_document`):
    {"content": "...", "title": "...", "format": "pdf"|"docx"|"both"}

The content may contain simple markdown:
    # Title
    ## Section
    - bullet
    plain paragraph
"""
from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

logger = logging.getLogger("jarvis.doc_tools")

_DOC_OUT_DIR = Path.home() / "Documents" / "Jarvis"


def _sanitize_filename(text: str) -> str:
    t = re.sub(r"[^A-Za-z0-9 _\-]+", "", text)[:60].strip()
    return t.replace(" ", "_") or "document"


def create_document(parameters: dict, player=None) -> str:
    """Create a styled PDF / DOCX from content.

    parameters:
        content : body text (plain or simple markdown)
        title   : document title (optional)
        author  : author name (optional)
        format  : "pdf" | "docx" | "both" (default: "pdf")
        path    : optional explicit output path (default: ~/Documents/Jarvis)
    """
    content = (parameters.get("content") or "").strip()
    if not content:
        return "What should I write? Give me the content, e.g. 'create a PDF that explains photosynthesis'."
    title = (parameters.get("title") or "").strip()
    author = (parameters.get("author") or "Jarvis").strip()
    fmt = (parameters.get("format") or "pdf").lower()
    fmt = "pdf" if "pdf" in fmt else ("docx" if "docx" in fmt else "both")

    out_dir = _DOC_OUT_DIR
    req_path = parameters.get("path")
    if req_path:
        p = Path(req_path).expanduser()
        if p.suffix.lower() in (".pdf", ".docx"):
            out_dir = p.parent
            if not title:
                title = p.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    base = _sanitize_filename(title or content[:40])
    today = date.today().isoformat()

    if not title:
        m = re.match(r"^#{1,2}\s+(.+)", content)
        title = m.group(1) if m else content[:50]

    results = []
    if fmt in ("pdf", "both"):
        results.append(_write_pdf(out_dir, base, title, author, today, content))
    if fmt in ("docx", "both"):
        results.append(_write_docx(out_dir, base, title, author, today, content))
    return " | ".join(results)


def _blocks(content: str):
    """Parse simple markdown into a list of (kind, text): heading / bullet / para."""
    kinds = []
    for line in content.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#"):
            level = min(len(s) - len(s.lstrip("#")), 3)
            kinds.append(("heading", s.lstrip("#").strip(), level))
        elif s.startswith(("- ", "* ", "+ ")):
            kinds.append(("bullet", s[2:].strip()))
        elif re.match(r"^\d+\.\s", s):
            kinds.append(("bullet", re.sub(r"^\d+\.\s+", "", s)))
        else:
            kinds.append(("para", s))
    return kinds


def _write_pdf(out_dir: Path, base: str, title: str, author: str,
               today: str, content: str) -> str:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Title page block
    pdf.set_font("Helvetica", "B", 22)
    pdf.multi_cell(0, 12, title, align="C")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(0, 7, f"Author: {author}", align="C", ln=1)
    pdf.cell(0, 7, f"Date: {today}", align="C", ln=1)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    blocks = _blocks(content)
    if blocks and blocks[0][0] == "heading" and blocks[0][1] == title:
        blocks = blocks[1:]  # title already rendered on the first page
    for kind, text, *rest in blocks:
        if kind == "heading":
            level = rest[0]
            size = 16 if level == 1 else (13 if level == 2 else 11)
            style = "B" if level <= 2 else "I"
            pdf.set_font("Helvetica", style, size)
            pdf.ln(4)
            pdf.multi_cell(0, 9, text)
            pdf.ln(1)
        elif kind == "bullet":
            pdf.set_font("Helvetica", "", 11)
            pdf.cell(6)
            pdf.multi_cell(0, 7, f"- {text}")
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, text)
            pdf.ln(2)

    path = out_dir / f"{base}.pdf"
    pdf.output(str(path))
    logger.info("[doc_tools] PDF written: %s", path)
    return f"PDF created: {path}"


def _write_docx(out_dir: Path, base: str, title: str, author: str,
                today: str, content: str) -> str:
    from docx import Document
    from docx.shared import Pt
    doc = Document()

    doc.add_heading(title, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Author: {author}   |   Date: {today}").italic = True
    meta.runs[0].font.size = Pt(9)

    blocks = _blocks(content)
    if blocks and blocks[0][0] == "heading" and blocks[0][1] == title:
        blocks = blocks[1:]  # title already rendered as the document heading
    for kind, text, *rest in blocks:
        if kind == "heading":
            level = rest[0]
            doc.add_heading(text, level=level)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    path = out_dir / f"{base}.docx"
    doc.save(str(path))
    logger.info("[doc_tools] DOCX written: %s", path)
    return f"Word document created: {path}"
